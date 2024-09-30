import os

import Voice_Record

from langchain_core.output_parsers import StrOutputParser, JsonOutputParser
from langchain_community.chat_models import ChatOllama
from langchain_community.vectorstores import Chroma
from langchain_community.chat_models import ChatOllama
from langchain_core.prompts import ChatPromptTemplate
from langchain_cohere import CohereEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_cohere import CohereRerank
from langchain.retrievers.contextual_compression import ContextualCompressionRetriever
from docx import Document
from langchain.docstore.document import Document as LC_Document
import speech_recognition as sr
import pyttsx3 

os.environ["COHERE_API_KEY"] = "XL0l1cfNDXPI4DFQf19IDQFy84hdeTZ3HSOVmnhs"
os.environ["LANGCHAIN_TRACING_V2"]="true"
os.environ["LANGCHAIN_ENDPOINT"]="https://api.smith.langchain.com"
os.environ["LANGCHAIN_API_KEY"]="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"
os.environ["LANGCHAIN_PROJECT"]="Metaverse Project"



print("Loading Agents ...")
llm = ChatOllama(model="llama3.1", format="json", temperature=0)
generation_llm = ChatOllama(model="llama3.1", temperature=2)
print("Agents loaded successfully")
print("Loading langchain doc")

# List of documents with corresponding URLs
# Process each document
docs = './Advisory.docx'

print("Formatting Information ...")

doc_path = './Advisory.docx'
doc = Document(doc_path)
full_text = "\n".join([para.text for para in doc.paragraphs])
text_splitter = RecursiveCharacterTextSplitter(chunk_size=1024, chunk_overlap=150)
doc_obj = [LC_Document(page_content=full_text)]
doc_splits = text_splitter.split_documents(doc_obj)
print("Documents Successfully Formatted")
print("Saving to ChromaDB ...")
vectorstore = Chroma.from_documents(
    documents=doc_splits,
    collection_name="rg-chroma",
    embedding= CohereEmbeddings(model="embed-english-v3.0")
)
print("ChromaDB successfully saved")
retriever = vectorstore.as_retriever(search_kwargs={"k": 15})
print("Retriever successfully configured")

def SpeakText(command):
    
    # Initialize the engine
    engine = pyttsx3.init()
    voices = engine.getProperty('voices')
    engine.setProperty('voice', voices[1].id)
    engine.say(command)
    engine.runAndWait()

def reranked_document(query):
    print("Retrieving related documents ...")
    compressor = CohereRerank(model="rerank-english-v3.0")
    compression_retriever = ContextualCompressionRetriever(base_compressor=compressor, base_retriever=retriever)
    print("Reranking Documents")
    compressed_docs = compression_retriever.invoke(query)
    print("Reranking Completed")
    return compressed_docs

def grade_harmfulness(query):

    harmfulness_grading_template="""system You are a grader assessing the harmfulness 
    of a generated content. 
    Analyze the following content for harmful elements. Specifically, check for any of the following:

    Violence: Does the content depict or encourage violence, including but not limited to physical harm, abuse, or threats?
    Hate Speech: Does the content contain any language that discriminates against, insults, or demeans individuals or groups based on race, religion, ethnicity, gender, sexual orientation, or other protected characteristics?
    Misinformation: Does the content spread false or misleading information, particularly related to health, safety, or critical social issues?
    Explicit Content: Does the content contain sexually explicit material, graphic descriptions of violence, or other adult themes inappropriate for general audiences?
    Mental Health Impact: Could the content potentially cause distress, anxiety, or other negative mental health effects to those who consume it?

    The goal is to filter out harmful contents. \n
    Give a binary score 'yes' or 'no' score to indicate whether the generated content is harmful. \n
    Provide the binary score as a JSON with a single key 'score' and no preamble or explanation.
    user
    Here is the generated content: \n\n {content} \n\n assistant
    """
    harmfulness_grading_prompt = ChatPromptTemplate.from_template(harmfulness_grading_template)

    harmfulness_grader_chain = ( harmfulness_grading_prompt | llm | JsonOutputParser())
    print("Grading Harmfulness ...")
    harmfulness_grade = harmfulness_grader_chain.invoke({"content": query})
    print("Grading Completed")

    return harmfulness_grade

def generate_answer(advisory, query):
    generation_template = """
    system
    You are Joy from Consult Anytime, a compassionate and understanding chatbot designed to support teenagers facing mental health challenges. 
    Your primary goal is to listen empathetically, validate their feelings, and provide gentle encouragement with advice to overcome their challenges. 
    Communicate in a warm, friendly, and non-judgmental manner. Avoid giving direct medical advice or diagnoses,
     but encourage users to seek support from trusted adults or professional mental health resources when appropriate.

    user
    Possible Advisory: {advisory}

    Teenager's Inquiry: {question}

    """

    generation_prompt = ChatPromptTemplate.from_template(generation_template)
    gen_chain = ( generation_prompt | generation_llm | StrOutputParser())
    print("Generating Response ...")
    generation = gen_chain.invoke({"advisory": advisory, "question": query})
    print("Response Generated")
    return generation



def check_score(data):
    return data.get('score') == 'yes'

r = sr.Recognizer()

SpeakText("Hi there, I am Joy, how are you today?")


query = Voice_Record.activate()

while query != "goodbye":
    answer_valid = False
    attempts = 0
    compressed_docs = reranked_document(query)
    while answer_valid == False and attempts < 2:
        generation = generate_answer(compressed_docs, query)
        harmfulness_grade = grade_harmfulness(generation)
        print(generation)
        if harmfulness_grade.get('score') == 'no':
            answer_valid = True
        attempts += 1
    if attempts == 2:
        SpeakText("Sorry, I am unable to answer you")
    else:
        SpeakText(generation)
    SpeakText("I hope it helped you, what else can I help you with?")
    query = Voice_Record.activate()



